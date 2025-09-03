import pandas as pd
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
import logging
import os

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class DocumentParser:
    def extract_text(self, file_path: str) -> str:
        """
        Extracts text from PDF, DOCX, TXT, and CSV files, including tables.
        """
        _, file_extension = os.path.splitext(file_path.lower())
        text = ""

        if file_extension == '.pdf':
            try:
                reader = PdfReader(file_path)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            except Exception as e:
                logging.error(f"Failed to extract text from PDF file: {e}")
        
        elif file_extension == '.docx':
            try:
                doc = DocxDocument(file_path)
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " | "
                        text += "\n"
            except Exception as e:
                logging.error(f"Failed to extract text from DOCX file: {e}")
        
        elif file_extension == '.txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            except Exception as e:
                logging.error(f"Failed to read TXT file: {e}")

        elif file_extension == '.csv':
            try:
                df = pd.read_csv(file_path, encoding='utf-8')
                text = df.to_string()
            except Exception as e:
                logging.error(f"Failed to read CSV file: {e}")
        
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        return text

    def parse_item_master(self, file_path: str) -> list:
        """
        Parses item master content from a CSV and standardizes column names.
        """
        try:
            df = pd.read_csv(file_path, encoding='utf-8')
            
            column_mapping = {
                'Part Number': 'part_number',
                'Item Code': 'part_number',
                'Description': 'description',
                'Item Name': 'material_name',
                'Proposed Registered Item Name': 'material_name',
                'Vendor': 'vendor_name',
                'UoM': 'uom',
                'Supplier': 'vendor_name'
            }
            
            standardized_df = df.rename(columns=column_mapping)
            
            standard_columns = ['material_name', 'part_number', 'description', 'vendor_name', 'uom']
            
            return standardized_df[[col for col in standard_columns if col in standardized_df.columns]].to_dict('records')
        except Exception as e:
            logging.error(f"Error parsing item master file: {e}")
            return []
     
